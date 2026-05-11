import json
import re
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
INPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "internship_diary_entries.json"
OUTPUT = Path(sys.argv[2]) if len(sys.argv) > 2 else ROOT / "Vaastav_Sanghvi_Internship_Diary.docx"


ACCENT = "1F4E79"
LIGHT = "EAF2F8"
MID = "D6EAF8"
TEXT = RGBColor(31, 41, 55)


def clean(value):
    if value is None:
        return ""
    text = str(value)
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def fmt_date(value, long=False):
    dt = datetime.strptime(value, "%Y-%m-%d")
    return dt.strftime("%d %B %Y" if long else "%d-%m-%Y")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.08
    run = para.add_run(clean(text) if text else "-")
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Aptos"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="B7C9D6", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_meta_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(11.5)
    set_table_borders(table)
    for label, value in pairs:
        row = table.add_row()
        row.cells[0].width = Cm(4.5)
        row.cells[1].width = Cm(11.5)
        set_cell_text(row.cells[0], label, bold=True, size=10, color=RGBColor(15, 76, 117))
        set_cell_text(row.cells[1], value, size=10)
        set_cell_shading(row.cells[0], LIGHT)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Aptos Display"
    run.font.color.rgb = RGBColor(15, 76, 117)
    run.font.bold = True
    return paragraph


def add_body_paragraph(doc, text):
    for chunk in clean(text).split("\n\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.08
        run = paragraph.add_run(chunk)
        run.font.name = "Aptos"
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT


def add_summary_table(doc, entries):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = ("Date", "Hours", "Internship", "Activity Summary")
    widths = (Cm(2.7), Cm(1.7), Cm(5.2), Cm(8.0))
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = widths[idx]
        set_cell_shading(cell, ACCENT)
        set_cell_text(cell, header, bold=True, size=8.5, color=RGBColor(255, 255, 255))
    for entry in entries:
        row = table.add_row()
        summary = clean(entry.get("description", "")).split("\n\n")[0]
        if len(summary) > 230:
            summary = summary[:227].rstrip() + "..."
        values = (
            fmt_date(entry["date"]),
            str(entry.get("hours") or "-"),
            clean(entry.get("internship", {}).get("name", "")),
            summary,
        )
        for idx, value in enumerate(values):
            row.cells[idx].width = widths[idx]
            set_cell_text(row.cells[idx], value, size=8)
    return table


def add_entry(doc, entry, number):
    add_heading(doc, f"Day {number}: {fmt_date(entry['date'], long=True)}", 2)
    add_meta_table(
        doc,
        [
            ("Internship", entry.get("internship", {}).get("name", "")),
            ("Hours Worked", f"{entry.get('hours') or '-'} hours"),
            ("Links / References", entry.get("links") or "-"),
            ("Blockers", entry.get("blockers") or "-"),
        ],
    )
    add_heading(doc, "Work Description", 3)
    add_body_paragraph(doc, entry.get("description") or "-")
    add_heading(doc, "Learning Outcomes", 3)
    add_body_paragraph(doc, entry.get("learnings") or "-")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = TEXT
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.color.rgb = RGBColor(15, 76, 117)
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 3"].font.size = Pt(10.5)


def add_footer(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Internship Diary")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(107, 114, 128)


def main():
    payload = json.loads(INPUT.read_text(encoding="utf-8-sig"))
    user = payload.get("user", {})
    entries = payload.get("entries", [])
    entries = sorted(entries, key=lambda item: item["date"])
    if not entries:
        raise SystemExit("No diary entries found.")

    total_hours = sum(float(item.get("hours") or 0) for item in entries)
    by_month = defaultdict(int)
    for item in entries:
        by_month[datetime.strptime(item["date"], "%Y-%m-%d").strftime("%B %Y")] += 1

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    add_footer(section)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("INTERNSHIP DIARY")
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 76, 117)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(clean(entries[0].get("internship", {}).get("name", "")))
    run.font.name = "Aptos"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(55, 65, 81)

    add_meta_table(
        doc,
        [
            ("Student Name", user.get("name", "")),
            ("Email", user.get("email", "")),
            ("Internship Period", f"{fmt_date(entries[0]['date'], long=True)} to {fmt_date(entries[-1]['date'], long=True)}"),
            ("Total Diary Entries", str(len(entries))),
            ("Total Hours Recorded", f"{total_hours:g} hours"),
            ("Source", "VTU Internyet student diary entries"),
            ("Generated On", datetime.now().strftime("%d %B %Y")),
        ],
    )

    add_heading(doc, "Monthly Record", 1)
    month_table = doc.add_table(rows=1, cols=2)
    set_table_borders(month_table)
    for idx, header in enumerate(("Month", "Entries")):
        set_cell_shading(month_table.rows[0].cells[idx], ACCENT)
        set_cell_text(month_table.rows[0].cells[idx], header, bold=True, color=RGBColor(255, 255, 255))
    for month, count in by_month.items():
        row = month_table.add_row()
        set_cell_text(row.cells[0], month, size=9)
        set_cell_text(row.cells[1], str(count), size=9)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1])
    add_heading(doc, "Diary Summary", 1)
    add_summary_table(doc, entries)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_footer(doc.sections[-1])
    add_heading(doc, "Daily Diary Entries", 1)
    for number, entry in enumerate(entries, start=1):
        add_entry(doc, entry, number)

    doc.core_properties.author = clean(user.get("name") or "Student")
    doc.core_properties.title = "Internship Diary"
    doc.core_properties.subject = "VTU Internyet internship diary entries"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
