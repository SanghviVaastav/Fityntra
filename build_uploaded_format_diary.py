import copy
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
TEMPLATE = Path(r"C:\Users\sejal\Downloads\Internship Dairy.docx")
INPUT = ROOT / "vaastav_current_diary_entries.json"
OUTPUT = ROOT / "Vaastav_Internship_Dairy_Format.docx"


def clean(value):
    if value is None:
        return ""
    text = str(value)
    replacements = {
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2013": "-",
        "\u2014": "-",
        "\u00a0": " ",
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    return re.sub(r"\s+", " ", text).strip()


def set_text_like(cell, text, reference_cell):
    cell.text = ""
    ref_para = reference_cell.paragraphs[0]
    para = cell.paragraphs[0]
    para.alignment = ref_para.alignment
    para.paragraph_format.left_indent = ref_para.paragraph_format.left_indent
    para.paragraph_format.right_indent = ref_para.paragraph_format.right_indent
    para.paragraph_format.first_line_indent = ref_para.paragraph_format.first_line_indent
    para.paragraph_format.space_before = ref_para.paragraph_format.space_before
    para.paragraph_format.space_after = ref_para.paragraph_format.space_after
    para.paragraph_format.line_spacing = ref_para.paragraph_format.line_spacing

    ref_run = ref_para.runs[0] if ref_para.runs else None
    run = para.add_run(text)
    if ref_run:
        run.bold = ref_run.bold
        run.italic = ref_run.italic
        run.underline = ref_run.underline
        run.font.name = ref_run.font.name
        run.font.size = ref_run.font.size
    else:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main():
    payload = json.loads(INPUT.read_text(encoding="utf-8-sig"))
    entries = sorted(payload["entries"], key=lambda item: item["date"])

    doc = Document(TEMPLATE)
    if not doc.tables:
        raise SystemExit("Template does not contain a table.")

    table = doc.tables[0]
    reference_row = table.rows[0]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    while len(table.rows) < len(entries):
        table.add_row()

    for row, entry in zip(table.rows, entries):
        date_text = datetime.strptime(entry["date"], "%Y-%m-%d").strftime("%d/%m/%Y")
        summary = clean(entry.get("description", ""))
        set_text_like(row.cells[0], date_text, reference_row.cells[0])
        set_text_like(row.cells[1], summary, reference_row.cells[1])

    doc.core_properties.title = "Internship Dairy"
    doc.core_properties.author = clean(payload.get("user", {}).get("name", ""))
    doc.save(OUTPUT)
    print(OUTPUT)
    print(len(entries))


if __name__ == "__main__":
    main()
